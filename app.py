# Updated: Force Heroku to use this app.py file
from flask import Flask, render_template, request, jsonify, send_from_directory
import os
import json
import sys
import importlib
import hashlib
import time
import logging
from datetime import datetime, date, timezone
from werkzeug.utils import secure_filename
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_wtf.csrf import CSRFProtect
from flask_talisman import Talisman
from sqlalchemy import text
import PyPDF2
from docx import Document
import io

# Database imports
from models import db, Resume, ContactInfo, ResumeSkill, Experience, Education, Project
from models import Certification, Achievement, MatchedSkill, MissingSkill, JobDescription
from models import ProcessingLog, Analytics
from config import get_config
from export_csv import register_export_listeners
from flask_migrate import Migrate
from logging_config import setup_logging

# Force reload of ats_components to get latest changes
if 'ats_components' in sys.modules:
    importlib.reload(sys.modules['ats_components'])

# Import our ATS components with fallback handling
try:
    from ats_components import ATSScorer, ResumeParser, ResumeExtractor, ContactExtractor, SemanticMatcher
    ML_AVAILABLE = True
except ImportError as e:
    logging.warning(f"ML dependencies not available: {e}. Using fallback implementations.")
    ML_AVAILABLE = False
    # Enhanced fallback classes with realistic scoring
    import re
    
    class ATSScorer:
        def __init__(self, *args, **kwargs):
            self.common_skills = [
                'python', 'java', 'javascript', 'react', 'angular', 'node.js', 'sql', 'mongodb',
                'aws', 'azure', 'docker', 'kubernetes', 'git', 'html', 'css', 'typescript',
                'spring', 'django', 'flask', 'machine learning', 'data science', 'pandas',
                'tensorflow', 'pytorch', 'linux', 'windows', 'agile', 'scrum', 'project management'
            ]
            
        def score_resume(self, resume_text, job_description=''):
            if not resume_text or len(resume_text.strip()) < 50:
                return {
                    'overall_score': 25,
                    'skills_score': 20,
                    'header_score': 30,
                    'experience_score': 25,
                    'projects_score': 20,
                    'education_score': 25,
                    'format_score': 30,
                    'matched_skills': [],
                    'missing_skills': [],
                    'message': 'Resume too short or invalid'
                }
                
            resume_lower = resume_text.lower()
            job_lower = job_description.lower() if job_description else ''
            
            # Enhanced Skills scoring with job matching
            found_skills = []
            job_skills = []
            
            # Extract skills from job description if provided
            if job_description and len(job_description) > 50:
                for skill in self.common_skills:
                    if skill.lower() in job_lower:
                        job_skills.append(skill)
            
            # Find skills in resume
            for skill in self.common_skills:
                if skill.lower() in resume_lower:
                    found_skills.append(skill)
            
            # Calculate base skills score
            base_skills = len(found_skills)
            skills_score = min(95, 25 + (base_skills * 6))  # Base 25, +6 per skill
            
            # Bonus for job-relevant skills
            if job_skills:
                relevant_matches = [skill for skill in found_skills if skill in job_skills]
                job_match_bonus = len(relevant_matches) * 8  # +8 per job-relevant skill
                skills_score = min(95, skills_score + job_match_bonus)
            
            # Header scoring (comprehensive contact info analysis)
            header_score = 20  # Lower base score
            email_found = bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text))
            phone_found = bool(re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', resume_text))
            
            if email_found:
                header_score += 30  # Email is crucial
            if phone_found:
                header_score += 20  # Phone is important
            if any(word in resume_lower for word in ['linkedin', 'github']):
                header_score += 15  # Professional profiles
            if any(word in resume_lower for word in ['portfolio', 'website', 'blog']):
                header_score += 10  # Additional online presence
            
            # Name detection bonus
            lines = resume_text.split('\n')
            potential_name_found = False
            for line in lines[:10]:  # Check first 10 lines
                if re.match(r'^[A-Z][a-zA-Z\s]{2,40}$', line.strip()):
                    potential_name_found = True
                    break
            if potential_name_found:
                header_score += 10
                
            header_score = min(95, header_score)
            
            # Enhanced Experience scoring
            experience_indicators = [
                'experience', 'worked', 'employed', 'position', 'role', 'job',
                'developed', 'managed', 'led', 'created', 'implemented', 'designed',
                'responsible', 'achieved', 'delivered', 'coordinated', 'supervised'
            ]
            
            exp_count = sum(1 for indicator in experience_indicators if indicator in resume_lower)
            date_patterns = len(re.findall(r'\b(19|20)\d{2}\b', resume_text))
            company_indicators = resume_lower.count('inc') + resume_lower.count('corp') + resume_lower.count('ltd')
            
            # Base experience score
            experience_score = 20  # Lower base
            experience_score += min(30, exp_count * 3)  # +3 per experience indicator, max 30
            experience_score += min(25, date_patterns * 4)  # +4 per year mentioned, max 25
            experience_score += min(20, company_indicators * 10)  # +10 per company indicator, max 20
            experience_score = min(95, experience_score)
            
            # Enhanced Projects scoring
            project_indicators = [
                'project', 'built', 'developed', 'created', 'designed', 'implemented',
                'github', 'repository', 'application', 'system', 'website', 'software'
            ]
            
            proj_count = sum(1 for indicator in project_indicators if indicator in resume_lower)
            github_mentions = resume_lower.count('github') + resume_lower.count('git')
            tech_diversity = len([skill for skill in found_skills if any(tech in skill.lower() for tech in ['python', 'java', 'react', 'node'])])
            
            projects_score = 15  # Lower base
            projects_score += min(35, proj_count * 4)  # +4 per project indicator, max 35
            projects_score += min(25, github_mentions * 12)  # +12 per git mention, max 25
            projects_score += min(25, tech_diversity * 5)  # +5 per diverse tech, max 25
            projects_score = min(95, projects_score)
            
            # Enhanced Education scoring
            edu_keywords = ['university', 'college', 'institute', 'school', 'academy']
            degree_keywords = ['bachelor', 'master', 'phd', 'degree', 'diploma', 'certificate', 'b.s', 'm.s', 'b.tech']
            
            edu_institutions = sum(1 for keyword in edu_keywords if keyword in resume_lower)
            degree_mentions = sum(1 for keyword in degree_keywords if keyword in resume_lower)
            gpa_mention = bool(re.search(r'\b(gpa|cgpa)\b|\b[3-4]\.[0-9]\b', resume_lower))
            
            education_score = 25  # Moderate base
            education_score += min(40, edu_institutions * 20)  # +20 per institution, max 40
            education_score += min(25, degree_mentions * 8)   # +8 per degree mention, max 25
            if gpa_mention:
                education_score += 15  # GPA mention bonus
            education_score = min(95, education_score)
            
            # Enhanced Format scoring
            word_count = len(resume_text.split())
            line_count = len([line for line in resume_text.split('\n') if line.strip()])
            char_count = len(resume_text)
            
            format_score = 30  # Base format score
            
            # Length appropriateness
            if 300 <= word_count <= 800:
                format_score += 25  # Ideal length
            elif 200 <= word_count < 300 or 800 < word_count <= 1200:
                format_score += 15  # Acceptable length
            elif word_count > 100:
                format_score += 10  # At least some content
            
            # Structure indicators
            bullet_points = resume_text.count('•') + resume_text.count('-') + resume_text.count('*')
            if bullet_points > 5:
                format_score += 15  # Good formatting
            
            # Section headers (case insensitive)
            sections = ['experience', 'education', 'skills', 'projects']
            section_count = sum(1 for section in sections if section in resume_lower)
            format_score += min(25, section_count * 8)  # +8 per section, max 25
            
            format_score = min(95, format_score)
            
            # Calculate weighted overall score
            overall_score = round(
                (skills_score * 0.35 + experience_score * 0.25 + projects_score * 0.15 + 
                 education_score * 0.12 + header_score * 0.08 + format_score * 0.05), 1
            )
            
            # Determine missing skills based on job description or common important skills
            missing_skills = []
            if job_skills:
                # Missing skills from job requirements
                missing_skills = [skill for skill in job_skills if skill not in found_skills]
            else:
                # Generic important skills not found
                important_skills = ['python', 'java', 'javascript', 'sql', 'git', 'html', 'css']
                missing_skills = [skill for skill in important_skills if skill not in found_skills]
            
            # Create performance message
            performance_level = "Excellent" if overall_score >= 85 else "Good" if overall_score >= 70 else "Fair" if overall_score >= 50 else "Needs Improvement"
            message = f'Enhanced fallback analysis: {performance_level} resume with {len(found_skills)} skills detected from {word_count} words'
            
            return {
                'overall_score': overall_score,
                'skills_score': skills_score,
                'header_score': header_score,
                'experience_score': experience_score,
                'projects_score': projects_score,
                'education_score': education_score,
                'format_score': format_score,
                'matched_skills': found_skills[:15],  # Top 15 matched skills
                'missing_skills': missing_skills[:8],  # Top 8 missing skills
                'message': message
            }
    
    class ResumeParser:
        def __init__(self, *args, **kwargs):
            self.common_skills = [
                'python', 'java', 'javascript', 'react', 'angular', 'node.js', 'sql', 'mongodb',
                'aws', 'azure', 'docker', 'kubernetes', 'git', 'html', 'css', 'typescript',
                'c++', 'c#', '.net', 'php', 'ruby', 'swift', 'kotlin', 'go', 'rust',
                'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn', 'machine learning',
                'data science', 'artificial intelligence', 'deep learning', 'nlp'
            ]
            
        def parse_resume(self, resume_text, *args, **kwargs):
            if not resume_text:
                return {'skills': [], 'experience': [], 'education': [], 'projects': [], 'certifications': []}
                
            resume_lower = resume_text.lower()
            lines = resume_text.split('\n')
            
            # Extract skills
            skills = [skill for skill in self.common_skills if skill.lower() in resume_lower]
            
            # Enhanced experience extraction
            experience = self._extract_experience(lines, resume_text)
            
            # Enhanced education extraction  
            education = self._extract_education(lines, resume_text)
            
            # Enhanced projects extraction
            projects = self._extract_projects(lines, resume_text)
            
            # Extract certifications
            certifications = self._extract_certifications(lines, resume_text)
                
            return {
                'skills': skills,
                'experience': experience,
                'education': education,
                'projects': projects,
                'certifications': certifications
            }
            
        def _extract_experience(self, lines, resume_text):
            """Extract work experience with proper formatting and detailed bullet points"""
            experience = []
            resume_lower = resume_text.lower()
            
            current_exp = None
            in_experience_section = False
            
            for i, line in enumerate(lines):
                line_clean = line.strip()
                line_lower = line_clean.lower()
                
                # Check if we're in experience section
                if 'experience' in line_lower and len(line_clean) < 30:
                    in_experience_section = True
                    continue
                    
                # Stop if we hit another section
                if any(section in line_lower for section in ['projects', 'education', 'skills', 'certifications']) and len(line_clean) < 30:
                    if current_exp:
                        experience.append(current_exp)
                        current_exp = None
                    in_experience_section = False
                    continue
                    
                if in_experience_section and line_clean:
                    # Look for company and title pattern (Company — Title)
                    if '—' in line_clean or ' - ' in line_clean:
                        if current_exp:
                            experience.append(current_exp)
                        
                        # Split company and title
                        if '—' in line_clean:
                            parts = line_clean.split('—', 1)
                        else:
                            parts = line_clean.split(' - ', 1)
                            
                        company = parts[0].strip()
                        title = parts[1].strip() if len(parts) > 1 else 'Position'
                        
                        current_exp = {
                            'company': company,
                            'title': title,
                            'description': '',
                            'dates': '',
                            'location': '',
                            'achievements': []
                        }
                    
                    # Look for date patterns
                    elif current_exp and re.search(r'\d{2}/\d{4}|\d{4}', line_clean):
                        current_exp['dates'] = line_clean
                    
                    # Look for location patterns
                    elif current_exp and ('|' in line_clean or 'remote' in line_lower or any(word in line_lower for word in ['delhi', 'mumbai', 'bangalore', 'india'])):
                        current_exp['location'] = line_clean
                    
                    # Look for bullet points and achievements
                    elif current_exp and line_clean.startswith(('•', '-', '*')):
                        achievement = line_clean.lstrip('•-* ').strip()
                        if len(achievement) > 10:  # Only meaningful achievements
                            current_exp['achievements'].append(achievement)
                    
                    # Add to description (paragraphs)
                    elif current_exp and len(line_clean) > 30 and not line_clean.startswith(('•', '-', '*')):
                        if current_exp['description']:
                            current_exp['description'] += ' '
                        current_exp['description'] += line_clean
            
            # Add final experience
            if current_exp:
                experience.append(current_exp)
            
            # Format for display with bullet points
            formatted_experience = []
            for exp in experience:
                # Main title line
                formatted_exp = f"{exp['company']} — {exp['title']}"
                
                # Add a bullet point with key information
                bullet_parts = []
                if exp.get('dates'):
                    bullet_parts.append(f"Duration: {exp['dates']}")
                if exp.get('location'):
                    bullet_parts.append(f"Location: {exp['location']}")
                
                # Use achievements if available, otherwise use description
                if exp.get('achievements'):
                    # Add first achievement as the main bullet point
                    main_achievement = exp['achievements'][0][:200]
                    if bullet_parts:
                        formatted_exp += f"\n• {' | '.join(bullet_parts)} | {main_achievement}"
                    else:
                        formatted_exp += f"\n• {main_achievement}"
                elif exp.get('description'):
                    # Use description as bullet point
                    desc_summary = exp['description'][:200]
                    if bullet_parts:
                        formatted_exp += f"\n• {' | '.join(bullet_parts)} | {desc_summary}"
                    else:
                        formatted_exp += f"\n• {desc_summary}"
                else:
                    # Default bullet point with just basic info
                    if bullet_parts:
                        formatted_exp += f"\n• {' | '.join(bullet_parts)}"
                    else:
                        formatted_exp += "\n• Professional experience in this role"
                    
                formatted_experience.append(formatted_exp)
            
            # Return structured data for JavaScript to format properly
            return experience[:5]
            
        def _extract_education(self, lines, resume_text):
            """Extract education with robust degree detection and proper formatting"""
            education = []
            resume_lower = resume_text.lower()
            
            current_edu = None
            in_education_section = False
            
            # Enhanced degree keywords for better detection
            degree_patterns = [
                r'bachelor\s+of\s+technology[\s–-]*.*',
                r'b\.?tech[\s–-]*.*',
                r'bachelor\s+of\s+science[\s–-]*.*',
                r'b\.?s\.?c?[\s–-]*.*',
                r'bachelor\s+of\s+arts[\s–-]*.*',
                r'b\.?a\.?[\s–-]*.*',
                r'bachelor\s+of\s+engineering[\s–-]*.*',
                r'b\.?e\.?[\s–-]*.*',
                r'master\s+of\s+technology[\s–-]*.*',
                r'm\.?tech[\s–-]*.*',
                r'master\s+of\s+science[\s–-]*.*',
                r'm\.?s\.?c?[\s–-]*.*',
                r'master\s+of\s+arts[\s–-]*.*',
                r'm\.?a\.?[\s–-]*.*',
                r'master\s+of\s+business\s+administration[\s–-]*.*',
                r'm\.?b\.?a\.?[\s–-]*.*',
                r'doctor\s+of\s+philosophy[\s–-]*.*',
                r'ph\.?d\.?[\s–-]*.*',
                r'diploma[\s–-]*.*',
                r'certificate[\s–-]*.*'
            ]
            
            for line in lines:
                line_clean = line.strip()
                line_lower = line_clean.lower()
                
                # Check if we're in education section
                if 'education' in line_lower and len(line_clean) < 30:
                    in_education_section = True
                    continue
                    
                # Stop at other sections
                if any(section in line_lower for section in ['experience', 'projects', 'skills', 'certifications']) and len(line_clean) < 30:
                    if current_edu:
                        education.append(current_edu)
                        current_edu = None
                    in_education_section = False
                    continue
                    
                if in_education_section and line_clean:
                    # Institution name
                    if any(word in line_lower for word in ['university', 'institute', 'college', 'technology', 'engineering', 'academy', 'school']):
                        if current_edu:
                            education.append(current_edu)
                        current_edu = {
                            'institution': line_clean,
                            'degree': '',
                            'dates': '',
                            'field': ''
                        }
                    
                    # Dates pattern
                    elif current_edu and re.search(r'\d{2}/\d{4}|\d{4}', line_clean):
                        current_edu['dates'] = line_clean
                    
                    # Enhanced degree detection using patterns
                    elif current_edu:
                        # Check for degree patterns in the line
                        degree_match = None
                        for pattern in degree_patterns:
                            match = re.search(pattern, line_lower)
                            if match:
                                degree_match = line_clean
                                break
                        
                        if degree_match and not current_edu['degree']:
                            current_edu['degree'] = degree_match
                        
                        # Fallback: if line contains degree-related keywords, capture it
                        elif not current_edu['degree'] and any(keyword in line_lower for keyword in ['bachelor', 'master', 'b.tech', 'btech', 'm.tech', 'mtech', 'b.sc', 'bsc', 'm.sc', 'msc', 'phd', 'ph.d', 'degree']):
                            current_edu['degree'] = line_clean
                    
                    # Field of study detection
                    elif current_edu and any(word in line_lower for word in ['computer science', 'engineering', 'science', 'arts', 'commerce', 'management']):
                        if not current_edu['field'] and 'degree' not in line_lower:
                            current_edu['field'] = line_clean
            
            # Add final education
            if current_edu:
                education.append(current_edu)
            
            # Return structured data with proper degree formatting
            for edu in education:
                # Ensure degree is properly formatted
                if not edu.get('degree') and not edu.get('field'):
                    edu['degree'] = 'Degree not specified'
                elif edu.get('degree') and edu.get('field') and edu['field'].lower() not in edu['degree'].lower():
                    edu['degree'] = f"{edu['degree']} - {edu['field']}"
            
            return education[:3]
            
        def _extract_projects(self, lines, resume_text):
            """Extract projects with detailed bullet points and achievements"""
            projects = []
            resume_lower = resume_text.lower()
            
            current_project = None
            in_projects_section = False
            
            for line in lines:
                line_clean = line.strip()
                line_lower = line_clean.lower()
                
                # Check if we're in projects section
                if 'project' in line_lower and len(line_clean) < 30:
                    in_projects_section = True
                    continue
                    
                # Stop at other sections
                if any(section in line_lower for section in ['experience', 'education', 'skills', 'certifications']) and len(line_clean) < 30:
                    if current_project:
                        projects.append(current_project)
                        current_project = None
                    in_projects_section = False
                    continue
                    
                if in_projects_section and line_clean:
                    # Project title with technologies (pattern: Title (Tech1, Tech2))
                    if '(' in line_clean and ')' in line_clean and not line_clean.startswith(('•', '-', '*')):
                        if current_project:
                            projects.append(current_project)
                        
                        # Extract title and technologies
                        title_part = line_clean.split('(')[0].strip()
                        tech_part = ''
                        if '(' in line_clean:
                            tech_part = line_clean[line_clean.find('(')+1:line_clean.find(')')]
                        
                        current_project = {
                            'name': title_part,
                            'technologies': tech_part,
                            'description': '',
                            'achievements': [],
                            'metrics': []
                        }
                    
                    # Project achievements/descriptions (bullet points)
                    elif current_project and line_clean.startswith(('•', '-', '*')):
                        achievement = line_clean.lstrip('•-* ').strip()
                        if len(achievement) > 15:  # Only meaningful achievements
                            current_project['achievements'].append(achievement)
                            
                            # Extract metrics/numbers from achievements
                            metrics = re.findall(r'\d+(?:\.\d+)?%|\d+(?:\.\d+)?\s*(?:k|million|billion)?', achievement.lower())
                            if metrics:
                                current_project['metrics'].extend(metrics[:2])  # Top 2 metrics
                    
                    # Additional description (non-bullet points)
                    elif current_project and len(line_clean) > 25 and not line_clean.startswith(('•', '-', '*')):
                        if current_project['description']:
                            current_project['description'] += ' '
                        current_project['description'] += line_clean
            
            # Add final project
            if current_project:
                projects.append(current_project)
            
            # Return structured data for JavaScript to format properly
            return projects[:5]
            
        def _extract_certifications(self, lines, resume_text):
            """Extract certifications with proper formatting"""
            certifications = []
            resume_lower = resume_text.lower()
            
            in_cert_section = False
            
            for line in lines:
                line_clean = line.strip()
                line_lower = line_clean.lower()
                
                # Check if we're in certifications section
                if 'certification' in line_lower and len(line_clean) < 30:
                    in_cert_section = True
                    continue
                    
                # Stop at other sections
                if any(section in line_lower for section in ['experience', 'education', 'skills', 'projects']) and len(line_clean) < 30:
                    in_cert_section = False
                    continue
                    
                if in_cert_section and line_clean and len(line_clean) > 10:
                    # Format: Certification Name (Year) or Certification Name
                    cert_text = line_clean
                    if line_clean.startswith(('•', '-', '*')):
                        cert_text = line_clean.lstrip('•-* ')
                    
                    # Extract year if present
                    year_match = re.search(r'\((\d{4})\)', cert_text)
                    if year_match:
                        cert_name = cert_text.replace(year_match.group(), '').strip()
                        year = year_match.group(1)
                        formatted_cert = f"{cert_name} ({year})"
                    else:
                        formatted_cert = cert_text
                    
                    certifications.append(formatted_cert)
            
            return certifications[:10]
    
    class ResumeExtractor:
        def __init__(self, *args, **kwargs):
            pass
            
        def extract_text(self, file_path):
            try:
                if file_path.endswith('.pdf'):
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text()
                        return text if text.strip() else "PDF text extraction failed - content may be image-based"
                        
                elif file_path.endswith('.docx'):
                    doc = Document(file_path)
                    return "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    
                elif file_path.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as file:
                        return file.read()
                        
                else:
                    return "Unsupported file format"
                    
            except Exception as e:
                return f"Text extraction error: {str(e)}"
    
    class ContactExtractor:
        def __init__(self, *args, **kwargs):
            pass
            
        def extract_contact_info(self, resume_text, *args, **kwargs):
            if not resume_text:
                return {'email': '', 'phone': '', 'name': ''}
                
            # Extract email
            email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text)
            email = email_match.group() if email_match else ''
            
            # Extract phone
            phone_match = re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', resume_text)
            phone = phone_match.group() if phone_match else ''
            
            # Extract name (basic heuristic - first few words that look like names)
            lines = resume_text.split('\n')
            name = ''
            for line in lines[:5]:  # Check first 5 lines
                line = line.strip()
                if len(line) > 0 and len(line) < 50:  # Reasonable name length
                    # Check if line contains mostly letters and spaces
                    if re.match(r'^[A-Za-z\s\.-]+$', line) and len(line.split()) <= 4:
                        name = line
                        break
            
            return {
                'email': email,
                'phone': phone,
                'name': name
            }
    
    class SemanticMatcher:
        def __init__(self, *args, **kwargs):
            pass
            
        def match_skills(self, resume_skills, job_skills, *args, **kwargs):
            # Simple string matching as fallback
            if not resume_skills or not job_skills:
                return [], job_skills[:5] if job_skills else []
                
            matched = []
            missing = []
            
            resume_lower = [skill.lower() for skill in resume_skills]
            
            for job_skill in job_skills:
                if job_skill.lower() in resume_lower:
                    matched.append(job_skill)
                else:
                    missing.append(job_skill)
                    
            return matched, missing[:10]  # Limit to 10 missing skills

# Initialize Flask app with configuration
app = Flask(__name__)
config = get_config()
app.config.from_object(config)

# Setup logging (must be done early)
app_logger, security_logger = setup_logging(app)

# Initialize security extensions
csrf = CSRFProtect(app)

# Configure Talisman for security headers
csp = {
    'default-src': "'self'",
    'script-src': [
        "'self'",
        "'unsafe-inline'",  # Required for inline scripts - consider removing in future
        "https://cdn.tailwindcss.com",
        "https://cdn.jsdelivr.net"
    ],
    'style-src': [
        "'self'",
        "'unsafe-inline'",  # Required for Tailwind CSS
        "https://cdn.tailwindcss.com",
        "https://cdnjs.cloudflare.com"
    ],
    'font-src': [
        "'self'",
        "https://cdnjs.cloudflare.com"
    ],
    'img-src': [
        "'self'",
        "data:",
        "https:"
    ]
}

# Only enable HTTPS enforcement in production
force_https = os.environ.get('FLASK_ENV') == 'production'
talisman = Talisman(
    app,
    force_https=force_https,
    strict_transport_security=True,
    content_security_policy=csp,
)

# Initialize database
db.init_app(app)

# Initialize Flask-Migrate for database migrations
migrate = Migrate(app, db)

# Register CSV export listeners so exports/all_resumes.csv is kept up-to-date
try:
    register_export_listeners(app, db, {
        'Resume': Resume,
        'ContactInfo': ContactInfo,
        'MatchedSkill': MatchedSkill,
        'MissingSkill': MissingSkill
    })
except Exception as e:
    print(f"Warning: could not register CSV export listeners: {e}")

# Rate limiter initialization (memory-based)

# Initialize metrics (optional, graceful degradation if not available)
try:
    from metrics import init_metrics, app_metrics
    metrics_available = True
except ImportError:
    metrics_available = False
    app_metrics = None
    print("Warning: Prometheus metrics not available. Install prometheus-flask-exporter for monitoring.")

# Configure rate limiter with memory backend
limiter = Limiter(
    key_func=get_remote_address,
    default_limits=["200 per hour", "50 per 15 minutes"],
    storage_uri="memory://"
)
app.logger.info("Rate limiter initialized with memory backend")

limiter.init_app(app)

# Initialize metrics if available
if metrics_available:
    try:
        prometheus_metrics = init_metrics(app)
        app.logger.info("Prometheus metrics enabled")
    except Exception as e:
        app.logger.warning(f"Metrics initialization failed: {e}")
        metrics_available = False

# Initialize database monitoring
try:
    from database_monitor import setup_database_monitoring, db_monitor, get_database_performance
    database_monitoring_available = True
    app.logger.info("Database monitoring initialized")
except ImportError:
    database_monitoring_available = False
    app.logger.warning("Database monitoring not available")

# Initialize database administration
try:
    from database_admin import initialize_database_admin, get_database_admin
    database_admin_available = True
    app.logger.info("Database administration initialized")
except ImportError:
    database_admin_available = False
    app.logger.warning("Database administration not available")

# Initialize backup system
try:
    from backup_manager import initialize_backup_manager, get_backup_manager
    backup_available = True
    app.logger.info("Backup system initialized")
except ImportError:
    backup_available = False
    app.logger.warning("Backup system not available")

# Initialize processing time tracking
try:
    from processing_tracker import processing_tracker, track_processing_time, get_processing_statistics
    processing_tracking_available = True
    app.logger.info("Processing time tracking initialized")
except ImportError:
    processing_tracking_available = False
    app.logger.warning("Processing time tracking not available")

# Initialize backup manager
try:
    from backup_manager import initialize_backup_manager, get_backup_manager
    backup_available = True
    app.logger.info("Backup manager initialized")
except ImportError:
    backup_available = False
    app.logger.warning("Backup manager not available")

# Setup database event listeners after db initialization
def setup_monitoring():
    """Setup monitoring systems after first request"""
    if database_monitoring_available:
        try:
            setup_database_monitoring(app, db)
            app.logger.info("Database monitoring event listeners registered")
        except Exception as e:
            app.logger.error(f"Failed to setup database monitoring: {e}")
    
    if database_admin_available:
        try:
            initialize_database_admin(db)
            app.logger.info("Database administration initialized")
        except Exception as e:
            app.logger.error(f"Failed to setup database administration: {e}")
    
    if backup_available:
        try:
            backup_dir = app.config.get('BACKUP_DIR', './backups')
            retention_days = app.config.get('BACKUP_RETENTION_DAYS', 30)
            initialize_backup_manager(db, backup_dir, retention_days)
            app.logger.info(f"Backup manager initialized (dir: {backup_dir}, retention: {retention_days} days)")
        except Exception as e:
            app.logger.error(f"Failed to setup backup manager: {e}")

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Create uploads directory if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def calculate_average_score():
    """Calculate average score from database"""
    try:
        avg_score = db.session.query(db.func.avg(Resume.overall_score)).scalar()
        return avg_score if avg_score else None
    except Exception as e:
        print(f"Error calculating average score: {e}")
        return None

def log_processing_stage(resume_id, stage, start_time, end_time=None, status="success", error_message=None):
    """Log processing stage for monitoring"""
    if not app.config.get('LOG_PROCESSING_TIMES', True):
        return
    
    try:
        duration = (end_time - start_time).total_seconds() if end_time else None
        
        log_entry = ProcessingLog(
            resume_id=resume_id,
            stage=stage,
            start_time=start_time,
            end_time=end_time,
            duration_seconds=duration,
            status=status,
            error_message=error_message
        )
        
        db.session.add(log_entry)
        db.session.commit()
    except Exception as e:
        print(f"Error logging processing stage: {e}")
        db.session.rollback()

def create_job_description_hash(job_description):
    """Create SHA256 hash of job description for deduplication"""
    return hashlib.sha256(job_description.encode('utf-8')).hexdigest()

def save_resume_data(resume_text, job_description, score_result, contact_info, parsed_sections, parsed_resume, filename, file_size, file_type, user_ip=None, session_id=None):
    """Save complete resume analysis to database"""
    try:
        start_time = datetime.utcnow()
        
        # Generate verdict based on analysis results
        classification, badge_color = classify_resume(score_result['overall_score'])
        verdict = generate_verdict(score_result, classification, parsed_resume, contact_info)
        
        # Create main resume record
        # Conditionally include header_job_title only if the DB table already has that column
        header_title_val = (parsed_sections.get('header_title') if isinstance(parsed_sections, dict) else None)
        include_header_col = False
        try:
            from sqlalchemy import inspect
            insp = inspect(db.engine)
            cols = [c['name'] for c in insp.get_columns('resumes')]
            include_header_col = 'header_job_title' in cols
        except Exception:
            include_header_col = False

        resume_kwargs = dict(
            filename=filename,
            file_size=file_size,
            file_type=file_type,
            overall_score=score_result['overall_score'],
            classification=classification,
            skills_score=score_result.get('skills_score', 0.0),
            header_score=score_result.get('header_score', 0.0),
            experience_score=score_result.get('experience_score', 0.0),
            projects_score=score_result.get('projects_score', 0.0),
            education_score=score_result.get('education_score', 0.0),
            format_score=score_result.get('format_score', 0.0),
            job_description_hash=create_job_description_hash(job_description),
            job_description_text=job_description,
            verdict=verdict,
            matched_skills_count=len(score_result.get('matched_skills', [])),
            missing_skills_count=len(score_result.get('missing_skills', [])),
            extracted_text=resume_text[:10000],  # Limit text length
            text_length=len(resume_text),
            user_ip=user_ip if app.config.get('STORE_IP_ADDRESSES', False) else None,
            session_id=session_id
        )

        if include_header_col and header_title_val:
            resume_kwargs['header_job_title'] = header_title_val[:300]

        resume = Resume(**resume_kwargs)
        
        db.session.add(resume)
        db.session.flush()  # Get the ID
        
        log_processing_stage(resume.id, 'save_main_record', start_time, datetime.utcnow())
        
        # Save contact information
        if contact_info:
            contact = ContactInfo(
                resume_id=resume.id,
                full_name=contact_info.get('name'),
                email=contact_info.get('email'),
                phone=contact_info.get('phone'),
                linkedin_url=contact_info.get('linkedin'),
                github_url=contact_info.get('github'),
                portfolio_url=contact_info.get('portfolio'),
                other_urls=contact_info.get('other_urls', {})
            )
            db.session.add(contact)
        
        # Save skills
        if parsed_sections.get('skills'):
            for skill in parsed_sections['skills'][:50]:  # Limit to 50 skills
                if isinstance(skill, str) and skill.strip():
                    skill_record = ResumeSkill(
                        resume_id=resume.id,
                        skill_name=skill[:200],  # Limit length
                        context_section='skills'
                    )
                    db.session.add(skill_record)
        
        # Save matched and missing skills
        for skill in score_result.get('matched_skills', [])[:50]:
            matched_skill = MatchedSkill(
                resume_id=resume.id,
                skill_name=skill[:200],
                match_type='exact'
            )
            db.session.add(matched_skill)
        
        for skill in score_result.get('missing_skills', [])[:50]:
            missing_skill = MissingSkill(
                resume_id=resume.id,
                skill_name=skill[:200],
                importance_level='required'
            )
            db.session.add(missing_skill)
        
        # Save experience entries (now formatted as strings)
        if parsed_sections.get('experience'):
            for exp in parsed_sections['experience'][:10]:  # Limit to 10 entries
                if isinstance(exp, str):
                    # Parse the formatted string to extract company and title
                    lines = exp.split('\n')
                    first_line = lines[0] if lines else exp
                    
                    # Extract company and title from "Company — Title" format
                    if '—' in first_line:
                        parts = first_line.split('—', 1)
                        company = parts[0].strip()
                        title = parts[1].strip()
                    else:
                        company = 'Professional Experience'
                        title = first_line.strip()
                    
                    description = '\n'.join(lines[1:]) if len(lines) > 1 else ''
                    
                    experience = Experience(
                        resume_id=resume.id,
                        company_name=company[:300],
                        job_title=title[:300],
                        description=description[:5000]
                    )
                    db.session.add(experience)
        
        # Save education entries (now formatted as strings)
        if parsed_sections.get('education'):
            for edu in parsed_sections['education'][:10]:  # Limit to 10 entries
                if isinstance(edu, str):
                    # Parse the formatted string
                    lines = edu.split('\n')
                    institution = lines[0] if lines else 'Educational Institution'
                    degree = lines[1] if len(lines) > 1 else 'Degree Program'
                    
                    education = Education(
                        resume_id=resume.id,
                        institution_name=institution[:300],
                        degree_name=degree[:300]
                    )
                    db.session.add(education)
        
        # Save projects (now formatted as strings)
        if parsed_sections.get('projects'):
            for proj in parsed_sections['projects'][:20]:  # Limit to 20 projects
                if isinstance(proj, str):
                    # Parse the formatted string
                    lines = proj.split('\n')
                    first_line = lines[0] if lines else proj
                    
                    # Extract project name and technologies from "Name (Tech1, Tech2)" format
                    if '(' in first_line and ')' in first_line:
                        name_part = first_line.split('(')[0].strip()
                        tech_part = first_line[first_line.find('(')+1:first_line.find(')')].strip()
                        technologies = [tech_part] if tech_part else []
                    else:
                        name_part = first_line.strip()
                        technologies = []
                    
                    description = '\n'.join(lines[1:]) if len(lines) > 1 else ''
                    
                    project = Project(
                        resume_id=resume.id,
                        project_name=name_part[:300],
                        description=description[:2000],
                        technologies_used=technologies
                    )
                    db.session.add(project)
        
        # Save certifications
        if parsed_sections.get('certifications'):
            for cert in parsed_sections['certifications'][:20]:  # Limit to 20 certs
                if isinstance(cert, dict):
                    certification = Certification(
                        resume_id=resume.id,
                        certification_name=cert.get('name', 'Certification')[:300],
                        issuing_organization=cert.get('issuer', 'Unknown')[:300],
                        issue_date=cert.get('date')
                    )
                    db.session.add(certification)
                elif isinstance(cert, str):
                    # Handle string certification entries
                    certification = Certification(
                        resume_id=resume.id,
                        certification_name=cert[:300],
                        issuing_organization='Unknown'
                    )
                    db.session.add(certification)
        
        # Save or update job description
        job_hash = create_job_description_hash(job_description)
        existing_job = JobDescription.query.filter_by(content_hash=job_hash).first()
        
        if existing_job:
            existing_job.usage_count += 1
            existing_job.last_used = datetime.utcnow()
        else:
            job_desc = JobDescription(
                content_hash=job_hash,
                description_text=job_description[:10000],  # Limit length
                usage_count=1
            )
            db.session.add(job_desc)
        
        # Commit all changes
        db.session.commit()
        
        log_processing_stage(resume.id, 'save_complete', start_time, datetime.utcnow())
        
        return resume.id
        
    except Exception as e:
        db.session.rollback()
        print(f"Error saving resume data: {e}")
        if 'resume' in locals():
            log_processing_stage(resume.id, 'save_complete', start_time, datetime.utcnow(), 
                               status='error', error_message=str(e))
        return None

# Initialize ATS components
ats_scorer = ATSScorer()
resume_parser = ResumeParser()
contact_extractor = ContactExtractor()

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_stream):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_stream):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_stream)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")

def extract_text_from_file(file):
    """Extract text from uploaded file based on extension"""
    filename = file.filename.lower()
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(io.BytesIO(file.read()))
    elif filename.endswith('.docx'):
        return extract_text_from_docx(io.BytesIO(file.read()))
    else:
        raise ValueError("Unsupported file format")

def classify_resume(overall_score):
    """Classify resume based on overall score"""
    # Updated thresholds to match the 59% rule-based system
    if overall_score < 50:
        return "No Fit", "danger"
    elif overall_score < 70:
        return "Potential Fit", "warning"
    else:
        return "Good Fit", "success"

def generate_verdict(score_result, classification, parsed_resume, contact_info):
    """Generate a comprehensive verdict for the resume analysis"""
    overall_score = score_result.get('overall_score', 0)
    matched_skills = score_result.get('matched_skills', [])
    missing_skills = score_result.get('missing_skills', [])
    
    verdict_parts = []
    
    # Overall assessment
    verdict_parts.append(f"Resume scored {overall_score}/100 and classified as '{classification}'.")
    
    # Skills analysis
    if matched_skills:
        verdict_parts.append(f"Strong skills match found: {', '.join(matched_skills[:3])}{'...' if len(matched_skills) > 3 else ''}.")
    
    if missing_skills:
        verdict_parts.append(f"Key skills to develop: {', '.join(missing_skills[:3])}{'...' if len(missing_skills) > 3 else ''}.")
    
    # Section-specific feedback
    sections = ['skills', 'experience', 'education', 'format', 'header', 'projects']
    low_sections = [s for s in sections if score_result.get(f'{s}_score', 0) < 60]
    
    if low_sections:
        verdict_parts.append(f"Areas for improvement: {', '.join(low_sections)}.")
    
    # Contact information
    if contact_info.get('email') and contact_info.get('phone'):
        verdict_parts.append("Complete contact information provided.")
    else:
        verdict_parts.append("Missing contact information (email or phone).")
    
    # Final recommendation
    if overall_score >= 70:
        verdict_parts.append("Recommended for further consideration.")
    elif overall_score >= 50:
        verdict_parts.append("Consider with additional screening.")
    else:
        verdict_parts.append("Not recommended for this position.")
    
    return " ".join(verdict_parts)

@app.route('/')
def index():
    """Render the main page"""
    return render_template('index.html')

@app.route('/about')
def about():
    """Render the about page"""
    return render_template('about.html')

@app.route('/ai')
def ai_models():
    """Render the AI models and technology page"""
    return render_template('ai.html')

@app.route('/analyze', methods=['POST'])
@limiter.limit("10 per minute")
@csrf.exempt
def analyze_resume():
    """Analyze uploaded resume"""
    # Generate session ID for processing tracking
    session_id = f"resume_{int(time.time() * 1000)}_{request.environ.get('REMOTE_ADDR', 'unknown')}"
    
    try:
        # Check if file is present
        if 'resume' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['resume']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Please upload PDF or DOCX'}), 400
        
        # Get file information for tracking
        filename = secure_filename(file.filename)
        file_type = filename.split('.')[-1].lower() if '.' in filename else 'unknown'
        
        # Read file to get size
        file_content = file.read()
        file_size = len(file_content)
        file.seek(0)  # Reset file pointer
        
        # Start processing session tracking
        if processing_tracking_available:
            processing_tracker.start_session(
                session_id=session_id,
                filename=filename,
                file_size=file_size,
                file_type=file_type
            )
        
        # Get job description
        job_description = request.form.get('job_description', '')
        
        if not job_description.strip():
            if processing_tracking_available:
                processing_tracker.end_session(session_id, success=False, error="Job description is required")
            return jsonify({'error': 'Job description is required'}), 400
        
        # Extract text from resume with tracking
        try:
            if processing_tracking_available:
                with processing_tracker.track_stage(session_id, "text_extraction", {"file_type": file_type, "file_size": file_size}):
                    resume_text = extract_text_from_file(file)
            else:
                resume_text = extract_text_from_file(file)
        except ValueError as e:
            if processing_tracking_available:
                processing_tracker.end_session(session_id, success=False, error=str(e))
            return jsonify({'error': str(e)}), 400
        
        if not resume_text.strip():
            error_msg = 'Could not extract text from file. File might be empty or corrupted.'
            if processing_tracking_available:
                processing_tracker.end_session(session_id, success=False, error=error_msg)
            return jsonify({'error': error_msg}), 400
        
        # Parse resume with tracking
        if processing_tracking_available:
            with processing_tracker.track_stage(session_id, "resume_parsing", {"text_length": len(resume_text)}):
                parsed_resume = resume_parser.parse_resume(resume_text)
        else:
            parsed_resume = resume_parser.parse_resume(resume_text)
        
        # Extract contact information with tracking
        if processing_tracking_available:
            with processing_tracker.track_stage(session_id, "contact_extraction"):
                contact_info = contact_extractor.extract_contact_info(resume_text)
        else:
            contact_info = contact_extractor.extract_contact_info(resume_text)
        
        # Score the resume with tracking
        if processing_tracking_available:
            with processing_tracker.track_stage(session_id, "resume_scoring", {"job_description_length": len(job_description)}):
                score_result = ats_scorer.score_resume(resume_text, job_description)
        else:
            score_result = ats_scorer.score_resume(resume_text, job_description)

        # Compute overall score as the integer average (truncate decimals) of the
        # six section percentages: education, experience, format, header,
        # projects, skills. This ensures the displayed ATS score matches the
        # requested behavior (e.g. 92.45 -> 92).
        section_scores = [
            score_result.get('education_score', 0),
            score_result.get('experience_score', 0),
            score_result.get('format_score', 0),
            score_result.get('header_score', 0),
            score_result.get('projects_score', 0),
            score_result.get('skills_score', 0),
        ]

        # Use integer truncation for the final ATS score (drop decimals)
        overall_score_int = int(sum(section_scores) / len(section_scores)) if section_scores else 0

        # Classify the resume using our computed overall score
        classification, badge_color = classify_resume(overall_score_int)

        # Calculate average score from database
        average_score = calculate_average_score()

        # Update overall score in score_result for database saving
        score_result['overall_score'] = overall_score_int

        # Save complete resume data to database with tracking
        user_ip = request.environ.get('HTTP_X_FORWARDED_FOR', request.environ.get('REMOTE_ADDR'))
        request_session_id = request.headers.get('X-Session-ID')
        
        if processing_tracking_available:
            with processing_tracker.track_stage(session_id, "database_save", {
                "overall_score": overall_score_int,
                "matched_skills": len(score_result.get('matched_skills', [])),
                "missing_skills": len(score_result.get('missing_skills', []))
            }):
                resume_id = save_resume_data(
                    resume_text=resume_text,
                    job_description=job_description,
                    score_result=score_result,
                    contact_info=contact_info,
                    parsed_sections=parsed_resume,
                    parsed_resume=parsed_resume,
                    filename=filename,
                    file_size=file_size,
                    file_type=file_type,
                    user_ip=user_ip,
                    session_id=request_session_id
                )
        else:
            resume_id = save_resume_data(
                resume_text=resume_text,
                job_description=job_description,
                score_result=score_result,
                contact_info=contact_info,
                parsed_sections=parsed_resume,
                parsed_resume=parsed_resume,
                filename=filename,
                file_size=file_size,
                file_type=file_type,
                user_ip=user_ip,
                session_id=request_session_id
            )
        
        # Prepare response
        response = {
            'success': True,
            'filename': filename,
            'classification': classification,
            'badge_color': badge_color,
            # Return overall score as integer (no decimals)
            'overall_score': overall_score_int,
            'average_score': round(average_score, 1) if average_score else None,
            'scores': {
                'skills': round(score_result['skills_score'], 1),
                'header': round(score_result['header_score'], 1),
                'experience': round(score_result['experience_score'], 1),
                'projects': round(score_result['projects_score'], 1),
                'education': round(score_result['education_score'], 1),
                'format': round(score_result['format_score'], 1)
            },
            'details': {
                'matched_skills': score_result.get('matched_skills', []),
                'missing_skills': score_result.get('missing_skills', []),
                'matched_count': len(score_result.get('matched_skills', [])),
                'missing_count': len(score_result.get('missing_skills', []))
            },
            'contact_info': contact_info,
            'parsed_sections': {
                'skills': parsed_resume.get('skills', []),
                'experience': parsed_resume.get('experience', []),
                'education': parsed_resume.get('education', []),
                'projects': parsed_resume.get('projects', []),
                'certifications': parsed_resume.get('certifications', [])
            },
            'resume_length': len(resume_text)
        }
        
        # End processing session successfully
        if processing_tracking_available:
            processing_tracker.end_session(session_id, success=True)
        
        return jsonify(response)
    
    except Exception as e:
        # End processing session with error
        if processing_tracking_available:
            processing_tracker.end_session(session_id, success=False, error=str(e))
        
        app.logger.error(f"Resume analysis error: {str(e)}")
        
        # Ensure we always return JSON, even on error
        response = jsonify({'error': f'An error occurred: {str(e)}'})
        response.headers['Content-Type'] = 'application/json'
        return response, 500

@app.route('/api/resumes')
def get_resumes():
    """Get paginated list of processed resumes"""
    try:
        page = request.args.get('page', 1, type=int)
        per_page = min(request.args.get('per_page', 20, type=int), 100)  # Max 100 per page
        
        # Query resumes with pagination
        resumes = Resume.query.order_by(Resume.upload_timestamp.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        
        resume_list = []
        for resume in resumes.items:
            resume_data = {
                'id': str(resume.id),
                'filename': resume.filename,
                'upload_timestamp': resume.upload_timestamp.isoformat(),
                'overall_score': resume.overall_score,
                'classification': resume.classification,
                'file_type': resume.file_type,
                'file_size': resume.file_size,
                'matched_skills_count': resume.matched_skills_count,
                'missing_skills_count': resume.missing_skills_count,
                'job_description': resume.job_description_text,
                'verdict': resume.verdict
            }
            resume_list.append(resume_data)
        
        return jsonify({
            'resumes': resume_list,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': resumes.total,
                'pages': resumes.pages,
                'has_next': resumes.has_next,
                'has_prev': resumes.has_prev
            }
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/resumes/<resume_id>')
def get_resume_details(resume_id):
    """Get detailed information about a specific resume"""
    try:
        resume = Resume.query.get_or_404(resume_id)
        
        # Get related data
        contact_info = ContactInfo.query.filter_by(resume_id=resume.id).first()
        skills = ResumeSkill.query.filter_by(resume_id=resume.id).all()
        experiences = Experience.query.filter_by(resume_id=resume.id).all()
        educations = Education.query.filter_by(resume_id=resume.id).all()
        projects = Project.query.filter_by(resume_id=resume.id).all()
        certifications = Certification.query.filter_by(resume_id=resume.id).all()
        matched_skills = MatchedSkill.query.filter_by(resume_id=resume.id).all()
        missing_skills = MissingSkill.query.filter_by(resume_id=resume.id).all()
        
        # Build detailed response
        response = {
            'resume': {
                'id': str(resume.id),
                'filename': resume.filename,
                'upload_timestamp': resume.upload_timestamp.isoformat(),
                'overall_score': resume.overall_score,
                'classification': resume.classification,
                'section_scores': {
                    'skills': resume.skills_score,
                    'header': resume.header_score,
                    'experience': resume.experience_score,
                    'projects': resume.projects_score,
                    'education': resume.education_score,
                    'format': resume.format_score
                },
                'file_info': {
                    'type': resume.file_type,
                    'size': resume.file_size,
                    'text_length': resume.text_length
                },
                'job_description': resume.job_description_text,
                'verdict': resume.verdict
            },
            'contact_info': {
                'full_name': contact_info.full_name if contact_info else None,
                'email': contact_info.email if contact_info else None,
                'phone': contact_info.phone if contact_info else None,
                'linkedin_url': contact_info.linkedin_url if contact_info else None,
                'github_url': contact_info.github_url if contact_info else None
            } if contact_info else None,
            'skills': [{'name': skill.skill_name, 'category': skill.skill_category} for skill in skills],
            'experiences': [{
                'company': exp.company_name,
                'title': exp.job_title,
                'start_date': exp.start_date.isoformat() if exp.start_date else None,
                'end_date': exp.end_date.isoformat() if exp.end_date else None,
                'is_current': exp.is_current,
                'location': exp.location
            } for exp in experiences],
            'education': [{
                'institution': edu.institution_name,
                'degree': edu.degree_name,
                'major': edu.major,
                'graduation_date': edu.graduation_date.isoformat() if edu.graduation_date else None
            } for edu in educations],
            'projects': [{
                'name': proj.project_name,
                'description': proj.description,
                'technologies': proj.technologies_used
            } for proj in projects],
            'certifications': [{
                'name': cert.certification_name,
                'issuer': cert.issuing_organization,
                'issue_date': cert.issue_date.isoformat() if cert.issue_date else None
            } for cert in certifications],
            'analysis': {
                'matched_skills': [skill.skill_name for skill in matched_skills],
                'missing_skills': [skill.skill_name for skill in missing_skills],
                'matched_count': len(matched_skills),
                'missing_count': len(missing_skills)
            }
        }
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/analytics/overview')
def get_analytics_overview():
    """Get analytics overview"""
    try:
        # Basic counts
        total_resumes = Resume.query.count()
        
        if total_resumes == 0:
            return jsonify({
                'total_resumes': 0,
                'message': 'No resume data available yet'
            })
        
        # Score statistics
        score_stats = db.session.query(
            db.func.avg(Resume.overall_score).label('avg_score'),
            db.func.min(Resume.overall_score).label('min_score'),
            db.func.max(Resume.overall_score).label('max_score')
        ).first()
        
        # Classification breakdown
        classifications = db.session.query(
            Resume.classification,
            db.func.count(Resume.id).label('count')
        ).group_by(Resume.classification).all()
        
        # Recent activity (last 30 days)
        from datetime import timedelta
        thirty_days_ago = datetime.utcnow() - timedelta(days=30)
        recent_count = Resume.query.filter(Resume.upload_timestamp >= thirty_days_ago).count()
        
        # Top skills (from matched skills)
        top_skills = db.session.query(
            MatchedSkill.skill_name,
            db.func.count(MatchedSkill.id).label('count')
        ).group_by(MatchedSkill.skill_name).order_by(
            db.func.count(MatchedSkill.id).desc()
        ).limit(10).all()
        
        # Most common missing skills
        top_missing = db.session.query(
            MissingSkill.skill_name,
            db.func.count(MissingSkill.id).label('count')
        ).group_by(MissingSkill.skill_name).order_by(
            db.func.count(MissingSkill.id).desc()
        ).limit(10).all()
        
        response = {
            'summary': {
                'total_resumes': total_resumes,
                'recent_resumes_30d': recent_count,
                'avg_score': round(float(score_stats.avg_score), 1) if score_stats.avg_score else 0,
                'score_range': {
                    'min': score_stats.min_score,
                    'max': score_stats.max_score
                }
            },
            'classifications': [
                {'classification': cls[0], 'count': cls[1], 'percentage': round(cls[1] / total_resumes * 100, 1)}
                for cls in classifications
            ],
            'top_skills': [
                {'skill': skill[0], 'count': skill[1]}
                for skill in top_skills
            ],
            'top_missing_skills': [
                {'skill': skill[0], 'count': skill[1]}
                for skill in top_missing
            ]
        }
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/analytics/trends')
def get_analytics_trends():
    """Get trending data over time"""
    try:
        # Daily resume count for last 30 days
        from datetime import timedelta
        thirty_days_ago = datetime.utcnow() - timedelta(days=30)
        
        # Build daily aggregates including classification breakdowns
        from sqlalchemy import case
        daily_counts = db.session.query(
            db.func.date(Resume.upload_timestamp).label('date'),
            db.func.count(Resume.id).label('count'),
            db.func.sum(case((Resume.classification == 'Good Fit', 1), else_=0)).label('good_fit'),
            db.func.sum(case((Resume.classification == 'Potential Fit', 1), else_=0)).label('potential_fit'),
            db.func.sum(case((Resume.classification == 'No Fit', 1), else_=0)).label('no_fit'),
            db.func.avg(Resume.overall_score).label('avg_score')
        ).filter(
            Resume.upload_timestamp >= thirty_days_ago
        ).group_by(
            db.func.date(Resume.upload_timestamp)
        ).order_by('date').all()

        trends = {
            'daily_activity': [],
            'classification_trends': []
        }

        for day in daily_counts:
            date_val = day.date.isoformat() if hasattr(day.date, 'isoformat') else str(day.date)
            trends['daily_activity'].append({
                'date': date_val,
                'resume_count': day.count,
                'avg_score': round(float(day.avg_score), 1) if day.avg_score else 0
            })

            trends['classification_trends'].append({
                'date': date_val,
                'good_fit': int(day.good_fit or 0),
                'potential_fit': int(day.potential_fit or 0),
                'no_fit': int(day.no_fit or 0)
            })
        
        return jsonify(trends)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/dashboard')
def dashboard():
    """Admin dashboard for viewing analytics"""
    return render_template('dashboard.html')

@app.route('/health')
def health():
    """Comprehensive health check endpoint for monitoring and load balancers"""
    health_checks = {}
    overall_healthy = True
    
    # Database connectivity check
    try:
        start_time = time.time()
        db.session.execute(text('SELECT 1')).scalar()
        db_time = time.time() - start_time
        
        health_checks['database'] = {
            'status': 'healthy',
            'response_time_ms': round(db_time * 1000, 2),
            'details': 'Database connection successful'
        }
    except Exception as e:
        health_checks['database'] = {
            'status': 'unhealthy',
            'error': str(e),
            'details': 'Database connection failed'
        }
        overall_healthy = False
        app.logger.error(f"Database health check failed: {str(e)}")
    
    # Cache status (memory-only)
    health_checks['cache'] = {
        'status': 'healthy',
        'response_time_ms': 0,
        'details': 'Memory-based caching and rate limiting active'
    }
    
    # File system checks
    directory_checks = {}
    required_dirs = ['uploads', 'exports', 'logs', 'instance']
    for directory in required_dirs:
        try:
            if not os.path.exists(directory):
                os.makedirs(directory, exist_ok=True)
            
            # Test write permissions
            test_file = os.path.join(directory, '.health_check')
            with open(test_file, 'w') as f:
                f.write('test')
            os.remove(test_file)
            
            directory_checks[directory] = {
                'status': 'healthy',
                'details': 'Directory accessible and writable'
            }
        except Exception as e:
            directory_checks[directory] = {
                'status': 'unhealthy',
                'error': str(e),
                'details': f'Directory {directory} not accessible or writable'
            }
            overall_healthy = False
    
    health_checks['filesystem'] = {
        'status': 'healthy' if all(d['status'] == 'healthy' for d in directory_checks.values()) else 'unhealthy',
        'directories': directory_checks
    }
    
    # ML Models check
    try:
        if ML_AVAILABLE:
            from model_manager import model_manager
            model_stats = model_manager.get_model_stats()
            
            health_checks['ml_models'] = {
                'status': 'healthy',
                'details': f"{model_stats['total_models']} models loaded",
                'memory_info': model_stats.get('memory_info', {}),
                'models': model_stats.get('models', {})
            }
        else:
            health_checks['ml_models'] = {
                'status': 'warning',
                'details': 'ML models disabled - running in lightweight mode',
                'memory_info': {},
                'models': {}
            }
    except Exception as e:
        health_checks['ml_models'] = {
            'status': 'degraded',
            'error': str(e),
            'details': 'ML models not accessible'
        }
    
    # Application metrics check
    if metrics_available and app_metrics:
        try:
            health_checks['metrics'] = {
                'status': 'healthy',
                'details': 'Prometheus metrics available'
            }
        except Exception as e:
            health_checks['metrics'] = {
                'status': 'degraded',
                'error': str(e),
                'details': 'Metrics collection issues'
            }
    else:
        health_checks['metrics'] = {
            'status': 'disabled',
            'details': 'Prometheus metrics not configured'
        }
    
    # Resource utilization check
    try:
        import psutil
        cpu_percent = psutil.cpu_percent(interval=0.1)
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('.')
        
        health_checks['resources'] = {
            'status': 'healthy' if cpu_percent < 90 and memory.percent < 90 and disk.percent < 90 else 'warning',
            'cpu_percent': cpu_percent,
            'memory_percent': memory.percent,
            'disk_percent': disk.percent,
            'details': f"CPU: {cpu_percent}%, Memory: {memory.percent}%, Disk: {disk.percent}%"
        }
        
        if cpu_percent > 95 or memory.percent > 95:
            health_checks['resources']['status'] = 'critical'
            
    except ImportError:
        health_checks['resources'] = {
            'status': 'unknown',
            'details': 'psutil not available for resource monitoring'
        }
    except Exception as e:
        health_checks['resources'] = {
            'status': 'error',
            'error': str(e),
            'details': 'Resource monitoring failed'
        }
    
    # Application-specific checks
    try:
        resume_count = Resume.query.count()
        health_checks['application'] = {
            'status': 'healthy',
            'resume_count': resume_count,
            'uptime': time.time() - app.config.get('START_TIME', time.time()),
            'details': f"Application running normally with {resume_count} resumes processed"
        }
    except Exception as e:
        health_checks['application'] = {
            'status': 'unhealthy',
            'error': str(e),
            'details': 'Application data access failed'
        }
        overall_healthy = False
    
    # Compile final health status
    overall_status = 'healthy' if overall_healthy else 'unhealthy'
    
    # Count status distribution
    status_counts = {}
    for check in health_checks.values():
        status = check.get('status', 'unknown')
        status_counts[status] = status_counts.get(status, 0) + 1
    
    health_data = {
        'status': overall_status,
        'service': 'ATS Resume Checker',
        'version': '1.0.0',
        'environment': app.config.get('ENV', 'unknown'),
        'timestamp': datetime.now(timezone.utc).isoformat(),
        'checks': health_checks,
        'summary': {
            'total_checks': len(health_checks),
            'status_distribution': status_counts,
            'critical_issues': sum(1 for c in health_checks.values() if c.get('status') in ['unhealthy', 'critical'])
        }
    }
    
    # Return appropriate HTTP status code
    if overall_status == 'healthy':
        status_code = 200
    elif any(c.get('status') == 'critical' for c in health_checks.values()):
        status_code = 503  # Service Unavailable
    else:
        status_code = 200  # OK but with warnings
    
    return jsonify(health_data), status_code

@app.route('/api/database/performance')
@limiter.limit("30 per minute")  # Limited access to performance data
def database_performance():
    """
    Database performance monitoring endpoint
    Returns query statistics, slow queries, and connection pool info
    """
    if not database_monitoring_available:
        return jsonify({
            'error': 'Database monitoring not available',
            'message': 'Database monitoring is not configured'
        }), 503
    
    try:
        performance_data = get_database_performance()
        
        # Add current connection pool status if available
        try:
            if hasattr(db.engine.pool, 'size'):
                pool_status = {
                    'pool_size': db.engine.pool.size(),
                    'checked_in': db.engine.pool.checkedin(),
                    'checked_out': db.engine.pool.checkedout(),
                    'overflow': db.engine.pool.overflow(),
                    'invalid': db.engine.pool.invalid()
                }
                performance_data['connection_pool'] = pool_status
        except Exception as e:
            app.logger.warning(f"Could not get connection pool status: {e}")
        
        # Add timestamp and metadata
        performance_data['timestamp'] = datetime.now(timezone.utc).isoformat()
        performance_data['monitoring_threshold'] = db_monitor.slow_query_threshold
        
        return jsonify(performance_data)
        
    except Exception as e:
        app.logger.error(f"Database performance monitoring error: {e}")
        return jsonify({
            'error': 'Failed to retrieve database performance data',
            'message': str(e)
        }), 500

@app.route('/api/database/slow-queries')
@limiter.limit("20 per minute")
def slow_queries():
    """
    Slow queries monitoring endpoint
    Returns detailed information about slow database queries
    """
    if not database_monitoring_available:
        return jsonify({
            'error': 'Database monitoring not available'
        }), 503
    
    try:
        limit = request.args.get('limit', 50, type=int)
        limit = min(limit, 100)  # Cap at 100 queries
        
        slow_queries_data = db_monitor.get_slow_queries(limit)
        
        return jsonify({
            'slow_queries': slow_queries_data,
            'count': len(slow_queries_data),
            'threshold': db_monitor.slow_query_threshold,
            'timestamp': datetime.now(timezone.utc).isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"Slow queries monitoring error: {e}")
        return jsonify({
            'error': 'Failed to retrieve slow queries data',
            'message': str(e)
        }), 500

@app.route('/api/database/query-stats')
@limiter.limit("20 per minute")
def query_statistics():
    """
    Query statistics endpoint
    Returns aggregated statistics about database query performance
    """
    if not database_monitoring_available:
        return jsonify({
            'error': 'Database monitoring not available'
        }), 503
    
    try:
        limit = request.args.get('limit', 20, type=int)
        limit = min(limit, 50)  # Cap at 50 query types
        
        query_stats = db_monitor.get_query_statistics(limit)
        summary = db_monitor.get_performance_summary()
        
        return jsonify({
            'query_statistics': query_stats,
            'summary': summary,
            'timestamp': datetime.now(timezone.utc).isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"Query statistics error: {e}")
        return jsonify({
            'error': 'Failed to retrieve query statistics',
            'message': str(e)
        }), 500

@app.route('/api/database/info')
@limiter.limit("10 per minute")
def database_info():
    """
    Database information endpoint
    Returns comprehensive database and system information
    """
    if not database_admin_available:
        return jsonify({
            'error': 'Database administration not available'
        }), 503
    
    try:
        admin = get_database_admin()
        if not admin:
            return jsonify({
                'error': 'Database admin not initialized'
            }), 503
            
        db_info = admin.get_database_info()
        connection_info = admin.get_connection_info()
        system_resources = admin.get_system_resources()
        
        return jsonify({
            'database': db_info,
            'connections': connection_info,
            'system': system_resources,
            'timestamp': datetime.now(timezone.utc).isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"Database info error: {e}")
        return jsonify({
            'error': 'Failed to retrieve database information',
            'message': str(e)
        }), 500

@app.route('/api/database/maintenance', methods=['POST'])
@limiter.limit("3 per hour")  # Very limited access to maintenance operations
def database_maintenance():
    """
    Database maintenance endpoint
    Run maintenance operations like ANALYZE, VACUUM, REINDEX
    """
    if not database_admin_available:
        return jsonify({
            'error': 'Database administration not available'
        }), 503
    
    try:
        admin = get_database_admin()
        if not admin:
            return jsonify({
                'error': 'Database admin not initialized'
            }), 503
        
        # Get operations from request
        data = request.get_json() or {}
        operations = data.get('operations', ['analyze'])
        
        # Validate operations
        valid_operations = ['analyze', 'vacuum', 'reindex']
        operations = [op for op in operations if op in valid_operations]
        
        if not operations:
            return jsonify({
                'error': 'No valid operations specified',
                'valid_operations': valid_operations
            }), 400
        
        app.logger.info(f"Running database maintenance operations: {operations}")
        results = admin.run_maintenance(operations)
        
        return jsonify(results)
        
    except Exception as e:
        app.logger.error(f"Database maintenance error: {e}")
        return jsonify({
            'error': 'Failed to run database maintenance',
            'message': str(e)
        }), 500

@app.route('/api/admin/dashboard')
@limiter.limit("20 per minute")
def admin_dashboard():
    """
    Administrative dashboard endpoint
    Returns comprehensive system status for admin interface
    """
    try:
        dashboard_data = {
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'service': 'ATS Resume Checker',
            'version': '1.0.0',
            'environment': app.config.get('ENV', 'unknown')
        }
        
        # Add database performance if available
        if database_monitoring_available:
            try:
                dashboard_data['database_performance'] = get_database_performance()
            except Exception as e:
                dashboard_data['database_performance'] = {'error': str(e)}
        
        # Add database info if available
        if database_admin_available:
            try:
                admin = get_database_admin()
                if admin:
                    dashboard_data['database_info'] = admin.get_database_info()
                    dashboard_data['system_resources'] = admin.get_system_resources()
            except Exception as e:
                dashboard_data['database_admin'] = {'error': str(e)}
        
        # Add application statistics
        try:
            from models import Resume, JobDescription, ProcessingLog
            dashboard_data['application_stats'] = {
                'total_resumes': Resume.query.count(),
                'total_job_descriptions': JobDescription.query.count(),
                'total_processing_logs': ProcessingLog.query.count()
            }
        except Exception as e:
            dashboard_data['application_stats'] = {'error': str(e)}
        
        # Add recent activity
        try:
            recent_resumes = Resume.query.order_by(Resume.created_at.desc()).limit(5).all()
            dashboard_data['recent_activity'] = [{
                'id': resume.id,
                'filename': resume.filename,
                'score': resume.overall_score,
                'created_at': resume.created_at.isoformat() if resume.created_at else None
            } for resume in recent_resumes]
        except Exception as e:
            dashboard_data['recent_activity'] = {'error': str(e)}
        
        return jsonify(dashboard_data)
        
    except Exception as e:
        app.logger.error(f"Admin dashboard error: {e}")
        return jsonify({
            'error': 'Failed to retrieve admin dashboard data',
            'message': str(e)
        }), 500

@app.route('/api/processing/statistics')
@limiter.limit("30 per minute")
def processing_statistics():
    """
    Processing time statistics endpoint
    Returns comprehensive processing performance data
    """
    if not processing_tracking_available:
        return jsonify({
            'error': 'Processing time tracking not available'
        }), 503
    
    try:
        stats = get_processing_statistics()
        return jsonify(stats)
        
    except Exception as e:
        app.logger.error(f"Processing statistics error: {e}")
        return jsonify({
            'error': 'Failed to retrieve processing statistics',
            'message': str(e)
        }), 500

@app.route('/api/processing/stages')
@limiter.limit("20 per minute")
def processing_stages():
    """
    Processing stages performance endpoint
    Returns detailed statistics for each processing stage
    """
    if not processing_tracking_available:
        return jsonify({
            'error': 'Processing time tracking not available'
        }), 503
    
    try:
        stage_name = request.args.get('stage')
        stats = processing_tracker.get_stage_statistics(stage_name)
        
        return jsonify({
            'stages': stats,
            'timestamp': datetime.now(timezone.utc).isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"Processing stages error: {e}")
        return jsonify({
            'error': 'Failed to retrieve processing stage statistics',
            'message': str(e)
        }), 500

@app.route('/api/processing/active')
@limiter.limit("60 per minute")
def active_processing():
    """
    Active processing sessions endpoint
    Returns currently running processing sessions
    """
    if not processing_tracking_available:
        return jsonify({
            'error': 'Processing time tracking not available'
        }), 503
    
    try:
        active_sessions = processing_tracker.get_active_sessions()
        
        return jsonify({
            'active_sessions': active_sessions,
            'count': len(active_sessions),
            'timestamp': datetime.now(timezone.utc).isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"Active processing error: {e}")
        return jsonify({
            'error': 'Failed to retrieve active processing sessions',
            'message': str(e)
        }), 500

@app.route('/api/processing/recent')
@limiter.limit("30 per minute")
def recent_processing():
    """
    Recent processing sessions endpoint
    Returns recently completed processing sessions
    """
    if not processing_tracking_available:
        return jsonify({
            'error': 'Processing time tracking not available'
        }), 503
    
    try:
        limit = request.args.get('limit', 50, type=int)
        limit = min(limit, 100)  # Cap at 100 sessions
        
        recent_sessions = processing_tracker.get_recent_sessions(limit)
        
        return jsonify({
            'recent_sessions': recent_sessions,
            'count': len(recent_sessions),
            'limit': limit,
            'timestamp': datetime.now(timezone.utc).isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"Recent processing error: {e}")
        return jsonify({
            'error': 'Failed to retrieve recent processing sessions',
            'message': str(e)
        }), 500

@app.route('/api/backup/create', methods=['POST'])
@limiter.limit("3 per hour")  # Very limited access to backup operations
def create_backup():
    """
    Create database backup endpoint
    Creates a new backup with optional custom name
    """
    if not backup_available:
        return jsonify({
            'error': 'Backup system not available'
        }), 503
    
    try:
        manager = get_backup_manager()
        if not manager:
            return jsonify({
                'error': 'Backup manager not initialized'
            }), 503
        
        data = request.get_json() or {}
        backup_name = data.get('backup_name')
        compress = data.get('compress', True)
        
        app.logger.info(f"Creating backup: {backup_name or 'auto-generated'}")
        result = manager.create_backup(backup_name, compress)
        
        if result.get('success'):
            return jsonify(result), 201
        else:
            return jsonify(result), 500
            
    except Exception as e:
        app.logger.error(f"Backup creation error: {e}")
        return jsonify({
            'error': 'Failed to create backup',
            'message': str(e)
        }), 500

@app.route('/api/backup/list')
@limiter.limit("20 per minute")
def list_backups():
    """
    List all backups endpoint
    Returns all available backups with metadata
    """
    if not backup_available:
        return jsonify({
            'error': 'Backup system not available'
        }), 503
    
    try:
        manager = get_backup_manager()
        if not manager:
            return jsonify({
                'error': 'Backup manager not initialized'
            }), 503
        
        backups = manager.list_backups()
        stats = manager.get_backup_statistics()
        
        return jsonify({
            'backups': backups,
            'statistics': stats,
            'count': len(backups),
            'timestamp': datetime.now(timezone.utc).isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"List backups error: {e}")
        return jsonify({
            'error': 'Failed to list backups',
            'message': str(e)
        }), 500

@app.route('/api/backup/restore/<backup_name>', methods=['POST'])
@limiter.limit("2 per hour")  # Very restricted restore operations
def restore_backup(backup_name):
    """
    Restore database from backup endpoint
    Restores database from specified backup
    """
    if not backup_available:
        return jsonify({
            'error': 'Backup system not available'
        }), 503
    
    try:
        manager = get_backup_manager()
        if not manager:
            return jsonify({
                'error': 'Backup manager not initialized'
            }), 503
        
        app.logger.warning(f"Attempting to restore from backup: {backup_name}")
        result = manager.restore_backup(backup_name)
        
        if result.get('success'):
            app.logger.info(f"Database restored successfully from backup: {backup_name}")
            return jsonify(result), 200
        else:
            app.logger.error(f"Backup restore failed: {result.get('error')}")
            return jsonify(result), 500
            
    except Exception as e:
        app.logger.error(f"Backup restore error: {e}")
        return jsonify({
            'error': 'Failed to restore backup',
            'message': str(e)
        }), 500

@app.route('/api/backup/delete/<backup_name>', methods=['DELETE'])
@limiter.limit("10 per hour")
def delete_backup(backup_name):
    """
    Delete backup endpoint
    Deletes specified backup file and metadata
    """
    if not backup_available:
        return jsonify({
            'error': 'Backup system not available'
        }), 503
    
    try:
        manager = get_backup_manager()
        if not manager:
            return jsonify({
                'error': 'Backup manager not initialized'
            }), 503
        
        app.logger.info(f"Deleting backup: {backup_name}")
        result = manager.delete_backup(backup_name)
        
        if result.get('success'):
            return jsonify(result), 200
        else:
            return jsonify(result), 404
            
    except Exception as e:
        app.logger.error(f"Backup deletion error: {e}")
        return jsonify({
            'error': 'Failed to delete backup',
            'message': str(e)
        }), 500

@app.route('/api/backup/statistics')
@limiter.limit("30 per minute")
def backup_statistics():
    """
    Backup statistics endpoint
    Returns backup system status and statistics
    """
    if not backup_available:
        return jsonify({
            'error': 'Backup system not available'
        }), 503
    
    try:
        manager = get_backup_manager()
        if not manager:
            return jsonify({
                'error': 'Backup manager not initialized'
            }), 503
        
        stats = manager.get_backup_statistics()
        return jsonify(stats)
        
    except Exception as e:
        app.logger.error(f"Backup statistics error: {e}")
        return jsonify({
            'error': 'Failed to retrieve backup statistics',
            'message': str(e)
        }), 500

# Error Handlers to ensure JSON responses
@app.errorhandler(400)
def bad_request(error):
    """Handle 400 errors with JSON response"""
    return jsonify({
        'error': 'Bad Request',
        'message': 'Invalid request data'
    }), 400

@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors with JSON response for API routes"""
    if request.path.startswith('/api/'):
        return jsonify({
            'error': 'Not Found',
            'message': 'The requested resource was not found'
        }), 404
    return render_template('index.html'), 404

@app.errorhandler(405)
def method_not_allowed(error):
    """Handle 405 errors with JSON response"""
    return jsonify({
        'error': 'Method Not Allowed',
        'message': 'The request method is not allowed for this endpoint'
    }), 405

@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle file too large errors with JSON response"""
    return jsonify({
        'error': 'File Too Large',
        'message': 'File size exceeds the maximum limit of 16MB'
    }), 413

@app.errorhandler(429)
def ratelimit_handler(error):
    """Handle rate limit errors with JSON response"""
    return jsonify({
        'error': 'Rate Limit Exceeded',
        'message': 'Too many requests. Please try again later.'
    }), 429

@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors with JSON response for API routes"""
    # Roll back any database changes
    db.session.rollback()
    
    if request.path.startswith('/api/') or request.path.startswith('/analyze'):
        return jsonify({
            'error': 'Internal Server Error',
            'message': 'An unexpected error occurred. Please try again.'
        }), 500
    return render_template('index.html'), 500

# Global error handler for unhandled exceptions
@app.errorhandler(Exception)
def handle_exception(error):
    """Handle all unhandled exceptions"""
    app.logger.error(f"Unhandled exception: {str(error)}", exc_info=True)
    
    # Roll back any database changes
    try:
        db.session.rollback()
    except:
        pass
    
    # Return JSON for API routes, HTML for others
    if request.path.startswith('/api/') or request.path.startswith('/analyze'):
        return jsonify({
            'error': 'Internal Server Error',
            'message': 'An unexpected error occurred. Please try again.'
        }), 500
    
    return render_template('index.html'), 500

if __name__ == '__main__':
    # Development server only - use wsgi.py for production
    print("🚀 Starting ATS Resume Checker Flask App (Development Mode)...")
    print("📝 Access the application at: http://localhost:5000")
    print("⚠️  For production, use: gunicorn wsgi:app")
    
    # Setup monitoring systems on startup
    with app.app_context():
        # Initialize database tables
        try:
            db.create_all()
            app.logger.info("Database tables initialized successfully")
        except Exception as e:
            app.logger.error(f"Database initialization error: {e}")
        
        setup_monitoring()
    
    # Get environment from config
    env = os.environ.get('FLASK_ENV', 'development')
    debug_mode = env == 'development'
    
    app.run(debug=debug_mode, host='0.0.0.0', port=5000)
